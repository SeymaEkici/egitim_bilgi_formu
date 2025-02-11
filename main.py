import os
import pandas as pd
from docx import Document
from thefuzz import fuzz, process  # Fuzzy string matching için

def find_best_match(key, choices, threshold=80):
    best_match, score = process.extractOne(key, choices, scorer=fuzz.token_sort_ratio)
    print(f"🔍 '{key}' -> '{best_match}' (%{score} benzerlik)")  # Debug için ekledik
    return best_match if score >= threshold else None

# .docx dosyasını okuma fonksiyonu
def read_docx(file_path):
    doc = Document(file_path)

    data = {
        "Eğitimin Adı": "",
        "Eğitmenin Adı-Soyadı": "",
        "Tahmini Eğitim Süresi": "",
        "Tahmini Eğitim Teslim Tarihi": "",
        "Eğitimin Amacı": "",
        "Eğitim Seviyesi": "",
        "Eğitim Niteliği": "",
        "Eğitim Gereksinimleri": "",
        "Hedef Kitle": "",
        "Eğitim Özeti": "",
        "Sıkça Sorulan Sorular": ""
    }

    # **1. Paragraflardan bilgileri çekelim**
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and ":" in text:  # ":" karakteri varsa işleme alıyoruz
            key, value = text.split(":", 1)
            key, value = key.strip(), value.strip()

            best_match = find_best_match(key, data.keys())
            if best_match:
                data[best_match] = value

    # **2. Tabloyu işleyelim (Ancak hücre sayısı kontrol etmiyoruz)**
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:  # Boş hücreleri atlıyoruz
                    print(f"📋 Hücredeki veri: {cell_text}")  # Debug: Hücre içeriği
                    # Hücredeki veriyi doğrudan işliyoruz
                    best_match = find_best_match(cell_text, data.keys())
                    if best_match:
                        data[best_match] = cell_text

    return data

# Veriyi CSV'ye kaydetme fonksiyonu
def save_to_csv(data, file_name):
    print("📊 Debug: Data'nın türü ve içeriği:")
    print(type(data))
    print(data)

    df = pd.DataFrame([data])
    print(f"📊 CSV veri:\n{df}")  # Debug: CSV'ye kaydedilecek veriyi göster

    df.to_csv(file_name, index=False)
    print(f"\n✔ CSV başarıyla oluşturuldu: {file_name}")


# HTML oluşturma fonksiyonu
def generate_html_from_csv(csv_path, template_path, output_dir):
    df = pd.read_csv(csv_path)

    # Jinja2 ortamını ayarla
    from jinja2 import Environment, FileSystemLoader
    env = Environment(loader=FileSystemLoader('templates'))
    template = env.get_template(template_path)

    os.makedirs(output_dir, exist_ok=True)

    for _, row in df.iterrows():
        data = {
            "egitim_adi": row.get("Eğitimin Adı", ""),
            "egitmen_adi": row.get("Eğitmenin Adı-Soyadı", ""),
            "egitim_suresi": row.get("Tahmini Eğitim Süresi", ""),
            "teslim_tarihi": row.get("Tahmini Eğitim Teslim Tarihi", ""),
            "egitim_amaci": row.get("Eğitimin Amacı", ""),
            "egitim_seviyesi": row.get("Eğitim Seviyesi", ""),
            "egitim_niteligi": row.get("Eğitim Niteliği", ""),
            "egitim_gereksinimleri": row.get("Eğitim Gereksinimleri", ""),
            "hedef_kitle": row.get("Hedef Kitle", ""),
            "egitim_ozeti": row.get("Eğitim Özeti", ""),
            "sikca_sorulan_sorular": row.get("Sıkça Sorulan Sorular", "")
        }

        egitim_adi = str(data['egitim_adi']) if data['egitim_adi'] else "default_egitim_adi"
        output_file = os.path.join(output_dir, f"{egitim_adi.replace(' ', '_')}.html")

        html_content = template.render(data)

        with open(output_file, "w", encoding="utf-8") as f:
            f.write(html_content)

        print(f"{output_file} oluşturuldu.")


if __name__ == "__main__":
    docx_path = input('Lütfen .docx dosyasının yolunu girin: ').strip('"')
    csv_path = "egitim_bilgileri.csv"
    template_path = "egitim_bilgi_formu.html"
    output_dir = "output"
    
    data = read_docx(docx_path)  # .docx dosyasından verileri al
    save_to_csv(data, csv_path)  # CSV'ye kaydet
    generate_html_from_csv(csv_path, template_path, output_dir)  # HTML oluştur
